from docx import Document;
from docx.shared import Inches, Pt,Cm;
import json;
import codecs
document = Document();

# set margin
margin = 2;
sections = document.sections
for section in sections:
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)
# read file
with codecs.open('file.txt', encoding='utf-8') as f:
    st = f.read();
# convert to json
js = json.loads(st); 
length = len(js['arrWord']);
# create table
table = document.add_table(length,2, 'Table Grid');
for cell in table.columns[0].cells:
    cell.width = Inches(4.0);
for cell in table.columns[1].cells:
    cell.width = Inches(7.0);
i = 0;
for row in table.rows:
    cell0 = row.cells[0];
    cell0.text = js['arrWord'][i] + ": " + js['arrTrans'][i];
    run = cell0.paragraphs[0].runs[0];
    run.font.bold = True;
    run.font.size = Pt(13);
    par = cell0.add_paragraph(js['arrIpa'][i]);
    par.space_before = Pt(0);
    row.cells[1].text = js['arrEx'][i].lstrip("\n");
    i+=1;

document.save('test.docx');
